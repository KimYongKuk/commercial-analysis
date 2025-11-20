"""
ë¬¸ì„œ ë¡œë” ë° ì²­í¬ ë¶„í•  ëª¨ë“ˆ

ë‹¤ì–‘í•œ í˜•ì‹ì˜ ë¬¸ì„œë¥¼ ì½ê³  RAGì— ì í•©í•œ í¬ê¸°ë¡œ ë¶„í• í•©ë‹ˆë‹¤.
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import os


class Document:
    """ë¬¸ì„œ ë°ì´í„° í´ë˜ìŠ¤"""

    def __init__(
        self,
        page_content: str,
        metadata: Optional[Dict[str, Any]] = None
    ):
        """
        Args:
            page_content: ë¬¸ì„œ ë‚´ìš©
            metadata: ë©”íƒ€ë°ì´í„° (íŒŒì¼ëª…, í˜ì´ì§€ ë²ˆí˜¸ ë“±)
        """
        self.page_content = page_content
        self.metadata = metadata or {}

    def __repr__(self):
        return f"Document(content='{self.page_content[:50]}...', metadata={self.metadata})"


class TextLoader:
    """í…ìŠ¤íŠ¸ íŒŒì¼ ë¡œë”"""

    def __init__(self, file_path: str, encoding: str = "utf-8"):
        """
        Args:
            file_path: íŒŒì¼ ê²½ë¡œ
            encoding: íŒŒì¼ ì¸ì½”ë”©
        """
        self.file_path = file_path
        self.encoding = encoding

    def load(self) -> List[Document]:
        """í…ìŠ¤íŠ¸ íŒŒì¼ ë¡œë“œ"""
        try:
            with open(self.file_path, "r", encoding=self.encoding) as f:
                content = f.read()

            metadata = {
                "source": os.path.basename(self.file_path),
                "file_path": self.file_path,
                "file_type": "txt"
            }

            return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            print(f"âŒ íŒŒì¼ ë¡œë“œ ì‹¤íŒ¨: {self.file_path}, ì—ëŸ¬: {e}")
            raise


class PDFLoader:
    """PDF íŒŒì¼ ë¡œë”"""

    def __init__(self, file_path: str):
        """
        Args:
            file_path: PDF íŒŒì¼ ê²½ë¡œ
        """
        self.file_path = file_path

    def load(self) -> List[Document]:
        """PDF íŒŒì¼ ë¡œë“œ"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(self.file_path)
            documents = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()

                if text.strip():  # ë¹ˆ í˜ì´ì§€ ì œì™¸
                    metadata = {
                        "source": os.path.basename(self.file_path),
                        "file_path": self.file_path,
                        "file_type": "pdf",
                        "page": page_num + 1,
                        "total_pages": len(reader.pages)
                    }
                    documents.append(Document(page_content=text, metadata=metadata))

            return documents
        except Exception as e:
            print(f"âŒ PDF ë¡œë“œ ì‹¤íŒ¨: {self.file_path}, ì—ëŸ¬: {e}")
            raise


class DOCXLoader:
    """DOCX íŒŒì¼ ë¡œë”"""

    def __init__(self, file_path: str):
        """
        Args:
            file_path: DOCX íŒŒì¼ ê²½ë¡œ
        """
        self.file_path = file_path

    def load(self) -> List[Document]:
        """DOCX íŒŒì¼ ë¡œë“œ"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(self.file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)

            metadata = {
                "source": os.path.basename(self.file_path),
                "file_path": self.file_path,
                "file_type": "docx"
            }

            return [Document(page_content=content, metadata=metadata)]
        except Exception as e:
            print(f"âŒ DOCX ë¡œë“œ ì‹¤íŒ¨: {self.file_path}, ì—ëŸ¬: {e}")
            raise


class DirectoryLoader:
    """ë””ë ‰í† ë¦¬ ë‚´ ëª¨ë“  ë¬¸ì„œ ë¡œë“œ"""

    def __init__(
        self,
        directory_path: str,
        glob_pattern: str = "**/*",
        supported_extensions: List[str] = None
    ):
        """
        Args:
            directory_path: ë””ë ‰í† ë¦¬ ê²½ë¡œ
            glob_pattern: íŒŒì¼ ê²€ìƒ‰ íŒ¨í„´
            supported_extensions: ì§€ì›í•˜ëŠ” í™•ì¥ì ë¦¬ìŠ¤íŠ¸
        """
        self.directory_path = Path(directory_path)
        self.glob_pattern = glob_pattern

        if supported_extensions is None:
            self.supported_extensions = [".txt", ".pdf", ".docx", ".md"]
        else:
            self.supported_extensions = supported_extensions

    def load(self) -> List[Document]:
        """ë””ë ‰í† ë¦¬ ë‚´ ëª¨ë“  ë¬¸ì„œ ë¡œë“œ"""
        documents = []

        for file_path in self.directory_path.glob(self.glob_pattern):
            if file_path.is_file() and file_path.suffix in self.supported_extensions:
                print(f"ğŸ“„ ë¡œë”© ì¤‘: {file_path.name}")
                try:
                    loader = self._get_loader(str(file_path))
                    docs = loader.load()
                    documents.extend(docs)
                except Exception as e:
                    print(f"âš ï¸  íŒŒì¼ ë¡œë“œ ì‹¤íŒ¨ (ê±´ë„ˆëœ€): {file_path.name}, ì—ëŸ¬: {e}")
                    continue

        print(f"âœ… ì´ {len(documents)}ê°œ ë¬¸ì„œ ë¡œë“œ ì™„ë£Œ")
        return documents

    def _get_loader(self, file_path: str):
        """íŒŒì¼ í™•ì¥ìì— ë”°ë¼ ì ì ˆí•œ ë¡œë” ë°˜í™˜"""
        ext = Path(file_path).suffix.lower()

        if ext == ".txt" or ext == ".md":
            return TextLoader(file_path)
        elif ext == ".pdf":
            return PDFLoader(file_path)
        elif ext == ".docx":
            return DOCXLoader(file_path)
        else:
            raise ValueError(f"ì§€ì›í•˜ì§€ ì•ŠëŠ” íŒŒì¼ í˜•ì‹: {ext}")


class TextSplitter:
    """í…ìŠ¤íŠ¸ ì²­í¬ ë¶„í• ê¸°"""

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 100,
        separator: str = "\n\n"
    ):
        """
        Args:
            chunk_size: ì²­í¬ í¬ê¸° (ë¬¸ì ìˆ˜)
            chunk_overlap: ì²­í¬ ê°„ ì˜¤ë²„ë© í¬ê¸°
            separator: êµ¬ë¶„ì (ë¬¸ë‹¨, ë¬¸ì¥ ë“±)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separator = separator

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """ë¬¸ì„œ ë¦¬ìŠ¤íŠ¸ë¥¼ ì²­í¬ë¡œ ë¶„í• """
        split_docs = []

        for doc in documents:
            chunks = self._split_text(doc.page_content)

            for i, chunk in enumerate(chunks):
                # ë©”íƒ€ë°ì´í„° ë³µì‚¬ ë° ì²­í¬ ì •ë³´ ì¶”ê°€
                chunk_metadata = doc.metadata.copy()
                chunk_metadata["chunk_index"] = i
                chunk_metadata["total_chunks"] = len(chunks)

                split_docs.append(
                    Document(page_content=chunk, metadata=chunk_metadata)
                )

        return split_docs

    def _split_text(self, text: str) -> List[str]:
        """í…ìŠ¤íŠ¸ë¥¼ ì²­í¬ë¡œ ë¶„í• """
        if not text or not text.strip():
            return []

        # êµ¬ë¶„ìë¡œ ë¨¼ì € ë¶„í• 
        splits = text.split(self.separator)
        chunks = []
        current_chunk = ""

        for split in splits:
            # í˜„ì¬ ì²­í¬ì— ì¶”ê°€í–ˆì„ ë•Œ í¬ê¸° í™•ì¸
            if len(current_chunk) + len(split) + len(self.separator) <= self.chunk_size:
                current_chunk += split + self.separator
            else:
                # í˜„ì¬ ì²­í¬ ì €ì¥
                if current_chunk.strip():
                    chunks.append(current_chunk.strip())

                # ìƒˆ ì²­í¬ ì‹œì‘ (ì˜¤ë²„ë© ê³ ë ¤)
                if self.chunk_overlap > 0 and current_chunk:
                    # ì´ì „ ì²­í¬ì˜ ëë¶€ë¶„ì„ í¬í•¨
                    overlap_text = current_chunk[-self.chunk_overlap:]
                    current_chunk = overlap_text + split + self.separator
                else:
                    current_chunk = split + self.separator

        # ë§ˆì§€ë§‰ ì²­í¬ ì €ì¥
        if current_chunk.strip():
            chunks.append(current_chunk.strip())

        # ë„ˆë¬´ ê¸´ ì²­í¬ëŠ” ê°•ì œë¡œ ë¶„í• 
        final_chunks = []
        for chunk in chunks:
            if len(chunk) > self.chunk_size * 1.5:  # 1.5ë°° ì´ˆê³¼ ì‹œ
                # ë¬¸ì ë‹¨ìœ„ë¡œ ê°•ì œ ë¶„í• 
                for i in range(0, len(chunk), self.chunk_size):
                    final_chunks.append(chunk[i:i + self.chunk_size])
            else:
                final_chunks.append(chunk)

        return final_chunks
